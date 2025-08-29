import pandas as pd
import requests
from bs4 import BeautifulSoup
import re
import os
import json
from docx import Document
from docx.shared import Inches, RGBColor
from docx.oxml.shared import OxmlElement, qn
from lxml import etree

def update_jama_attachments(url, session, project_ID, attachment_ID, file_path, basic_oauth):
    """
    Updates attachments in the description field of Jama items based on a list
    of Global IDs from an Excel file.

    Args:
        url (str): The base URL of the Jama Connect instance.
        session (requests.Session): The authenticated session object.
        project_ID (int): The ID of the project containing the items.
        attachment_ID (int): The item type ID for attachments.
        file_path (str): The file path to the Excel file with Global IDs.
        basic_oauth (str): The authentication method ('basic' or 'oauth').
    """
    try:
        # Ensure the URL is clean and doesn't have a trailing slash
        clean_url = url.rstrip('/')

        # Use the correct API path based on the working curl command
        api_path = 'rest/v1'

        # Create the DOWNLOADED folder if it doesn't exist
        download_folder = "DOWNLOADED"
        os.makedirs(download_folder, exist_ok=True)

        # Step 1: Get all items in the specified project using pagination
        print("Fetching all items from the specified project...")
        all_items_data = []
        start_at = 0
        max_results = 50  # Set a reasonable page size
        
        while True:
            # Corrected URL: use the /items endpoint with the 'project' query parameter
            project_items_url = f'{clean_url}/{api_path}/items?project={project_ID}&startAt={start_at}&maxResults={max_results}'
            response = session.get(project_items_url)
            response.raise_for_status()
            
            items_page = response.json().get('data', [])
            all_items_data.extend(items_page)
            
            # Check for pagination: break the loop if no more items are returned
            if len(items_page) < max_results:
                break
            
            start_at += max_results

        # Create a list of item API IDs and Global IDs
        all_items = {item['fields']['globalId']: item['id'] for item in all_items_data}
        
        # Step 2: Read the Excel file and create the map
        print("Reading the Excel file and creating the map...")
        df = pd.read_excel(file_path)
        # Corrected line to use the new column name
        global_ids_to_update = df['Attribute Value'].tolist()

        item_map = {}
        for global_id in global_ids_to_update:
            if global_id in all_items:
                item_map[all_items[global_id]] = global_id
            else:
                print(f"Warning: Global ID '{global_id}' not found in the project. Skipping.")

        if not item_map:
            print("No matching items found to update. Exiting.")
            return

        # Step 3: Loop through the map
        enumeration = 1
        processed_items = []
        for api_id, global_id in item_map.items():
            try:
                print(f"\nProcessing item with API ID: {api_id} and Global ID: {global_id}")

                # Scrape the description rich text field for the source HTML
                item_url = f'{clean_url}/{api_path}/items/{api_id}'
                response = session.get(item_url)
                response.raise_for_status()
                item_data = response.json().get('data', {})
                rich_text_html = item_data['fields'].get('description', '')

                if not rich_text_html:
                    print("No rich text content found in the description field. Skipping.")
                    processed_items.append({'ID': global_id, 'URL': f'{clean_url}/perspective.req#/items/{api_id}?projectId={project_ID}', 'Status': 'Skipped'})
                    continue

                # Find the source HTML pointing to the pasted image
                soup = BeautifulSoup(rich_text_html, 'html.parser')
                img_tag = soup.find('img')
                if not img_tag or 'src' not in img_tag.attrs:
                    print("No image found in the rich text field. Skipping.")
                    processed_items.append({'ID': global_id, 'URL': f'{clean_url}/perspective.req#/items/{api_id}?projectId={project_ID}', 'Status': 'Skipped'})
                    continue

                image_download_url = img_tag['src']
                
                # Use a regex to extract the abstract item ID from the URL
                match = re.search(r'attachment/(\d+)', image_download_url)
                if not match:
                    print("Could not find abstract item ID in the image URL. Skipping.")
                    processed_items.append({'ID': global_id, 'URL': f'{clean_url}/perspective.req#/items/{api_id}?projectId={project_ID}', 'Status': 'Skipped'})
                    continue
                
                abstract_item_id = match.group(1)

                # Attempt to download the file content from the official API endpoint
                print(f"Downloading old attachment with ID: {abstract_item_id}")
                file_download_url = f'{clean_url}/{api_path}/attachments/{abstract_item_id}/file'
                
                try:
                    file_response = session.get(file_download_url, stream=True)
                    file_response.raise_for_status()

                    # Process the downloaded file content
                    content_disposition = file_response.headers.get('Content-Disposition', '')
                    filename_list = re.findall(r'filename="([^"]+)"', content_disposition)
                    
                    if filename_list:
                        filename = filename_list[0]
                    else:
                        filename = os.path.basename(image_download_url)
                        print(f"Warning: Filename not found in header. Using URL as fallback: {filename}")

                    temp_file_path = os.path.join(download_folder, filename)
                    with open(temp_file_path, 'wb') as f:
                        for chunk in file_response.iter_content(chunk_size=8192):
                            f.write(chunk)
                            
                except requests.exceptions.HTTPError as http_err:
                    if http_err.response.status_code == 404:
                        print(f"Official API download failed (404). Falling back to direct URL download.")
                        
                        # Fallback to direct URL download using a new session
                        direct_download_url = f"{clean_url}/{image_download_url.lstrip('/')}"
                        file_response = session.get(direct_download_url, stream=True)
                        file_response.raise_for_status()
                        
                        filename = os.path.basename(direct_download_url)
                        temp_file_path = os.path.join(download_folder, filename)
                        with open(temp_file_path, 'wb') as f:
                            for chunk in file_response.iter_content(chunk_size=8192):
                                f.write(chunk)
                    else:
                        raise http_err

                # Rename the attachment
                name, extension = os.path.splitext(filename)
                new_filename = f"{name}_{enumeration:03d}{extension}"

                # Step A: Create a new attachment item (metadata)
                print(f"Creating new attachment item for: {new_filename}")
                create_attachment_url = f'{clean_url}/{api_path}/projects/{project_ID}/attachments'
                metadata_payload = {
                    'fields': {
                        'name': new_filename,
                        'description': 'Attachment renamed and re-uploaded via API script.'
                    }
                }
                
                headers = {
                    'Content-Type': 'application/json',
                    'Accept': 'application/json'
                }
                create_response = session.post(create_attachment_url, data=json.dumps(metadata_payload), headers=headers)
                create_response.raise_for_status()
                new_attachment_id = create_response.json()['meta']['id']
                print(f"Attachment item created with ID: {new_attachment_id}")

                # Step B: Upload the renamed file content to the newly created attachment item
                print(f"Uploading file content for: {new_filename}")
                upload_file_url = f'{clean_url}/{api_path}/attachments/{new_attachment_id}/file'
                
                with open(temp_file_path, 'rb') as f:
                    upload_response = session.put(upload_file_url, data=f, headers=session.headers)
                    upload_response.raise_for_status()
                
                print("File content successfully uploaded.")

                # Step C: Update the HTML in the description field
                print("Updating the item's description field...")
                new_attachment_url = f'{clean_url}/attachment/{new_attachment_id}/{new_filename}'
                new_html = rich_text_html.replace(image_download_url, new_attachment_url)
                
                update_url = f'{clean_url}/{api_path}/items/{api_id}'
                update_payload = {'fields': {'description': new_html}}
                update_response = session.put(update_url, json=update_payload)
                update_response.raise_for_status()
                
                print(f"Successfully updated item {global_id} with new attachment ID: {new_attachment_id}")
                os.remove(temp_file_path)
                
                processed_items.append({'ID': global_id, 'URL': f'{clean_url}/perspective.req#/items/{api_id}?projectId={project_ID}', 'Status': 'Updated'})
                enumeration += 1

            except requests.exceptions.HTTPError as http_err:
                print(f"HTTP error for item {global_id}: {http_err}. Skipping.")
                processed_items.append({'ID': global_id, 'URL': f'{clean_url}/perspective.req#/items/{api_id}?projectId={project_ID}', 'Status': 'Skipped'})
            except Exception as e:
                print(f"An unexpected error occurred while processing item {global_id}: {e}. Skipping.")
                processed_items.append({'ID': global_id, 'URL': f'{clean_url}/perspective.req#/items/{api_id}?projectId={project_ID}', 'Status': 'Skipped'})

        # Step 4: Create the Word document
        doc = Document()
        doc.add_heading('Jama Connect Item Update Report', level=1)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        # Add headers
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Global ID'
        hdr_cells[1].text = 'Item URL'
        hdr_cells[2].text = 'Status'
        
        for item in processed_items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item['ID'])
            
            # Create a new paragraph inside the cell for the hyperlink
            p = row_cells[1].paragraphs[0]
            # Updated function call with the correct URL format
            add_hyperlink(p, item['URL'], "Link")
            
            row_cells[2].text = item['Status']
        
        report_filename = "Jama_Update_Report.docx"
        doc.save(report_filename)
        
        print("\nOperation completed successfully.")
        print(f"Created a report file named: {report_filename}")
        print(f"File path: {os.path.abspath(report_filename)}")

    except requests.exceptions.HTTPError as http_err:
        print(f"A network or authentication error occurred: {http_err}")
    except Exception as e:
        print(f"An unexpected error occurred: {e}")

def add_hyperlink(paragraph, url, text, underline=True):
    """
    Adds a hyperlink to a paragraph without setting a custom color.
    This bypasses potential compatibility issues with RGBColor on some systems.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)
    
    new_run.append(rPr)
    new_run.text = text
    
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink